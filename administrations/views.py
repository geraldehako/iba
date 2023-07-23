from django.shortcuts import render
from administrations.models import Notes,Services,Postes,Catalogues, HistoriquesVers,Avis
from administrations.models import Etudiants, Matieres, Anneescolaires,Statutannees, Moyennes, Decoupages,Documents,Missions,Genres
from administrations.models import Ecoles,Contacturgences,Personnels,Autorisationsauditeur,Autorisationsenseignant,Autorisationspersonnel
from administrations.models import Salles,Typeecoles,Typematieres,Niveaux,DossierScolaires,Catalogues,HistoriquesAbs,DossierFinances
from accounts.models import Utilisateur
from .forms import EtudiantRegistration, MoyennemRegistration,TypematRegistration,MatiereRegistration,SessionRegistration,TypeinstRegistration,MisRegistration
from .forms import EcoleRegistration,EtudiantmRegistration,CatalogueRegistration,HistoriqueabsRegistration,EtudiantclasseRegistration,AutopersRegistration
from .forms import EnseignantRegistration, AnneescolaireRegistration,DocumentForm,SalleRegistration,SallemRegistration,NiveauRegistration,AutoensRegistration,AutoaudRegistration
from .forms import AvisRegistration,PersonnelRegistration,MessageForm
from accounts.forms import UserCreationForm
from .models import Enseignants
from django.conf import settings
from django.core.serializers import serialize
from django.shortcuts import redirect, get_object_or_404, redirect
from django.contrib import messages
from django.urls import reverse
from datetime import datetime, timedelta
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from PIL import Image
from django.core.files.storage import FileSystemStorage
from django.http import FileResponse, HttpResponse
from django.core.cache import cache
from django.contrib.auth.hashers import make_password
from django.forms import inlineformset_factory
import uuid
import docx # type: ignore
from django.views.generic.edit import CreateView
from django.urls import reverse_lazy
from django.http import JsonResponse
from django.db.models import Count, Q,Sum
import matplotlib.pyplot as plt
from io import BytesIO
import base64
from datetime import date

@login_required(login_url='/accounts/')
def menusuper(request):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    salle = Salles.objects.filter(Anneescolaire=statuta)
    num_salle = Salles.objects.filter(Anneescolaire=statuta).count()
    num_students = Etudiants.objects.filter(Salle__in=salle).count()

    context = {
        'num_salle': num_salle,
        'num_students': num_students,
        'annee': statuta,
    }  

    return render(request, 'Pages/menusuper.html', context)

@login_required()
def menu(request):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    salle = Salles.objects.filter(Anneescolaire=statuta)
    num_salle = Salles.objects.filter(Anneescolaire=statuta).count()
    num_students = Etudiants.objects.filter(Salle__in=salle).count()

    context = {
        'num_salle': num_salle,
        'num_students': num_students,
        'annee': statuta,
    }  

    return render(request, 'Pages/menu.html', context)

# Create your views here.
# DASHBOARD...................................................................................................................................
@login_required(login_url='/accounts/')
def dashboard(request):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    statuti = Decoupages.objects.get(Statutdecoup='1')
    salle = Salles.objects.filter(Anneescolaire=statuta)
    num_salle = Salles.objects.filter(Anneescolaire=statuta).count()
    num_students = Etudiants.objects.filter(Salle__in=salle).count()

    # DONUT ETUDIANT----------------------------------------------------------------------------
    # Récupérer les décomptes du nombre d'étudiants de chaque genre depuis la base de données
    male_count = Etudiants.objects.filter(Sexe='1').count()
    female_count = Etudiants.objects.filter(Sexe='2').count()

    # DONUT ENSEIGNANT----------------------------------------------------------------------------
    # Récupérer les décomptes du nombre d'étudiants de chaque genre depuis la base de données
    male_countens = Enseignants.objects.filter(Sexe='1').count()
    female_countens = Enseignants.objects.filter(Sexe='2').count()

    # BAR ETUDIANT PAR ECOLE----------------------------------------------------------------------------
    # Obtenez le nombre d'étudiants par école pour l'année scolaire active
    etudiants_par_ecole = Salles.objects.filter(Anneescolaire=statuta).values('Etablissement').annotate(num_etudiants=Count('etudiants'))

    # Récupérez les noms des écoles et le nombre d'étudiants
    ecoles = [Ecoles.objects.get(id=item['Etablissement']).Etablissement for item in etudiants_par_ecole] # type: ignore
    nombre_etudiants = [item['num_etudiants'] for item in etudiants_par_ecole]

    # Récupérer les objets HistoriquesAbs depuis la base de données
    historiques_abs = HistoriquesAbs.objects.all()

    # Initialiser des listes pour stocker les dates de cours, les heures d'absence et les heures de présence
    dates_cours = []
    heures_absence = []
    heures_presence = []

    # Itérer sur les objets HistoriquesAbs pour extraire les dates de cours, les heures d'absence et les heures de présence
    for historique in historiques_abs:
        date_cours = historique.Datecours
        heure_debut = historique.Heuredebut
        heure_fin = historique.Heurefin

        # Calculer les heures d'absence et les heures de présence
        duree_cours = heure_fin.hour - heure_debut.hour
        duree_absence = duree_cours  # Remplacez cette logique par la logique réelle de calcul des heures d'absence
        duree_presence = duree_cours - duree_absence

        # Ajouter les données à leurs listes respectives
        dates_cours.append(date_cours)
        heures_absence.append(duree_absence)
        heures_presence.append(duree_presence)

    # Récupérer l'année scolaire active    
    anneescolaire_active = Anneescolaires.objects.get(Statutannee='1')

    # Obtenir le nombre de dossiers financiers par classe pour l'année scolaire active
    dossier_counts = DossierFinances.objects.filter(Anneescolaire=anneescolaire_active).values('Salle__Salle').annotate(count=Count('id'))

    bar2_labels = []
    bar2_data = []

    for dossier_count in dossier_counts:
        salle = dossier_count['Salle__Salle']
        count = dossier_count['count']

        bar2_labels.append(salle)
        bar2_data.append(count)

    context = {
        'num_salle': num_salle,
        'num_students': num_students,
        'annee': statuta,
        'male_count': male_count,
        'female_count': female_count,
        'male_countens': male_countens,
        'female_countens': female_countens,
        'ecoles': ecoles,
        'nombre_etudiants': nombre_etudiants,
        'dates_cours': dates_cours,
        'heures_absence': heures_absence,
        'heures_presence': heures_presence,
        'bar2_labels': bar2_labels,
        'bar2_data': bar2_data,
    }


    return render(request, 'Pages/dashboard.html', context)


# DASHBOARD...................................................................................................................................
@login_required(login_url='/accounts/')
def dashboardacad(request):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    statuti = Decoupages.objects.get(Statutdecoup='1')
    salle = Salles.objects.filter(Anneescolaire=statuta)
    num_salle = Salles.objects.filter(Anneescolaire=statuta).count()
    num_students = Etudiants.objects.filter(Salle__in=salle).count()

    # DONUT ETUDIANT----------------------------------------------------------------------------
    # Récupérer les décomptes du nombre d'étudiants de chaque genre depuis la base de données
    male_count = Etudiants.objects.filter(Sexe='1').count()
    female_count = Etudiants.objects.filter(Sexe='2').count()

    # DONUT ENSEIGNANT----------------------------------------------------------------------------
    # Récupérer les décomptes du nombre d'étudiants de chaque genre depuis la base de données
    male_countens = Enseignants.objects.filter(Sexe='1').count()
    female_countens = Enseignants.objects.filter(Sexe='2').count()

    # BAR ETUDIANT PAR ECOLE----------------------------------------------------------------------------
    # Obtenez le nombre d'étudiants par école pour l'année scolaire active
    etudiants_par_ecole = Salles.objects.filter(Anneescolaire=statuta).values('Etablissement').annotate(num_etudiants=Count('etudiants'))

    # Récupérez les noms des écoles et le nombre d'étudiants
    ecoles = [Ecoles.objects.get(id=item['Etablissement']).Etablissement for item in etudiants_par_ecole]
    nombre_etudiants = [item['num_etudiants'] for item in etudiants_par_ecole]

    # Récupérer les objets HistoriquesAbs depuis la base de données
    historiques_abs = HistoriquesAbs.objects.all()

    # Initialiser des listes pour stocker les dates de cours, les heures d'absence et les heures de présence
    dates_cours = []
    heures_absence = []
    heures_presence = []

    # Itérer sur les objets HistoriquesAbs pour extraire les dates de cours, les heures d'absence et les heures de présence
    for historique in historiques_abs:
        date_cours = historique.Datecours
        heure_debut = historique.Heuredebut
        heure_fin = historique.Heurefin

        # Calculer les heures d'absence et les heures de présence
        duree_cours = heure_fin.hour - heure_debut.hour
        duree_absence = duree_cours  # Remplacez cette logique par la logique réelle de calcul des heures d'absence
        duree_presence = duree_cours - duree_absence

        # Ajouter les données à leurs listes respectives
        dates_cours.append(date_cours)
        heures_absence.append(duree_absence)
        heures_presence.append(duree_presence)

    # Récupérer l'année scolaire active    
    anneescolaire_active = Anneescolaires.objects.get(Statutannee='1')

    # Obtenir le nombre de dossiers financiers par classe pour l'année scolaire active
    dossier_counts = DossierFinances.objects.filter(Anneescolaire=anneescolaire_active).values('Salle__Salle').annotate(count=Count('id'))

    bar2_labels = []
    bar2_data = []

    for dossier_count in dossier_counts:
        salle = dossier_count['Salle__Salle']
        count = dossier_count['count']

        bar2_labels.append(salle)
        bar2_data.append(count)

    context = {
        'num_salle': num_salle,
        'num_students': num_students,
        'annee': statuta,
        'male_count': male_count,
        'female_count': female_count,
        'male_countens': male_countens,
        'female_countens': female_countens,
        'ecoles': ecoles,
        'nombre_etudiants': nombre_etudiants,
        'dates_cours': dates_cours,
        'heures_absence': heures_absence,
        'heures_presence': heures_presence,
        'bar2_labels': bar2_labels,
        'bar2_data': bar2_data,
    }


    return render(request, 'Pages/dashboardacademique.html', context)

@login_required(login_url='/accounts/')
def dashboardcaisse(request):
    today = datetime.now().date()
    start_date = today - timedelta(days=30)  # Période de 30 jours

    historiques = HistoriquesVers.objects.filter(Datevers__range=[start_date, today])

    # Agréger les données par jour
    aggregated_data = historiques.values('Datevers').annotate(
        count=Count('id'),
        total_amount=Sum('Montantvers')
    )

    bar_labels = []
    bar_data_count = []
    bar_data_amount = []

    for data in aggregated_data:
        date = data['Datevers'].strftime('%Y-%m-%d')
        bar_labels.append(date)
        bar_data_count.append(data['count'])
        bar_data_amount.append(data['total_amount'])

    # Récupérer l'année scolaire active
    anneescolaire_active = Anneescolaires.objects.get(Statutannee='1')

    # Obtenir le nombre de dossiers financiers par classe pour l'année scolaire active
    dossier_counts = DossierFinances.objects.filter(Anneescolaire=anneescolaire_active).values('Salle__Salle').annotate(count=Count('id'))

    bar2_labels = []
    bar2_data = []

    for dossier_count in dossier_counts:
        salle = dossier_count['Salle__Salle']
        count = dossier_count['count']

        bar2_labels.append(salle)
        bar2_data.append(count)

    # BAR ETUDIANT PAR ECOLE----------------------------------------------------------------------------
    # Obtenez le nombre d'étudiants par école pour l'année scolaire active
    etudiants_par_ecole = Salles.objects.filter(Anneescolaire=anneescolaire_active).values('Etablissement').annotate(num_etudiants=Count('etudiants'))

    # Récupérez les noms des écoles et le nombre d'étudiants
    ecoles = [Ecoles.objects.get(id=item['Etablissement']).Etablissement for item in etudiants_par_ecole]
    nombre_etudiants = [item['num_etudiants'] for item in etudiants_par_ecole]

    context = {
        'bar_labels': bar_labels,
        'bar_data_count': bar_data_count,
        'bar_data_amount': bar_data_amount,
        'bar2_labels': bar2_labels,
        'bar2_data': bar2_data,
        'ecoles': ecoles,
        'annee': anneescolaire_active,
        'nombre_etudiants': nombre_etudiants
    }

    return render(request, 'Pages/dashboardcaisse.html', context)


@login_required()
def dashboardrh(request):

    # DONUT ETUDIANT----------------------------------------------------------------------------
    # Récupérer les décomptes du nombre d'étudiants de chaque genre depuis la base de données
    male_count = Etudiants.objects.filter(Sexe='1').count()
    female_count = Etudiants.objects.filter(Sexe='2').count()

    # DONUT ENSEIGNANT----------------------------------------------------------------------------
    # Récupérer les décomptes du nombre d'étudiants de chaque genre depuis la base de données
    male_countens = Enseignants.objects.filter(Sexe='1').count()
    female_countens = Enseignants.objects.filter(Sexe='2').count()

    today = datetime.now().date()
    start_date = today - timedelta(days=30)  # Période de 30 jours

    # Récupérer l'année scolaire active
    annee_scolaire_active = Anneescolaires.objects.get(Statutannee='1')

    # Obtenir les données sur le nombre d'absences par mois
    absence_counts = HistoriquesAbs.objects.filter(Anneescolaire=annee_scolaire_active, Datecours__range=[start_date, today]).values('Datecours__month').annotate(count=Count('id'))

    # Préparer les données pour le graphique
    labels = []
    data = []

    for absence_count in absence_counts:
        month_number = absence_count['Datecours__month']
        count = absence_count['count']

        # Récupérer le nom du mois correspondant au numéro de mois
        month_name = datetime(1900, month_number, 1).strftime('%B')

        labels.append(month_name)
        data.append(count)

    # Obtenir la date actuelle
    today = date.today()

    # Récupérer les missions en cours
    missions_encours = Missions.objects.filter(Datefin__gt=today).count()
    autper_encours = Autorisationspersonnel.objects.filter(Datefin__gt=today).count()
    autens_encours = Autorisationsenseignant.objects.filter(Datefin__gt=today).count()
    autaud_encours = Autorisationsauditeur.objects.filter(Datefin__gt=today).count()
    activeauditeur = Utilisateur.objects.filter(is_active=1,role="ETUDIANT").count()
    activeens = Utilisateur.objects.filter(is_active=1,role="ENSEIGNANT").count()

    # Obtenez les dates de naissance des personnels, étudiants et enseignants
    dates_naissance_personnels = Personnels.objects.values_list('Datenaissance', flat=True)
    dates_naissance_etudiants = Etudiants.objects.values_list('Datenaissance', flat=True)
    dates_naissance_enseignants = Enseignants.objects.values_list('Datenaissance', flat=True)

    # Obtenez l'année actuelle
    annee_actuelle = date.today().year

    # Calculez les âges pour chaque groupe
    ages_personnels = [annee_actuelle - date_naissance.year for date_naissance in dates_naissance_personnels]
    ages_etudiants = [annee_actuelle - date_naissance.year for date_naissance in dates_naissance_etudiants]
    ages_enseignants = [annee_actuelle - date_naissance.year for date_naissance in dates_naissance_enseignants]

    # Comptez le nombre d'individus dans chaque tranche d'âge pour chaque groupe
    tranches_age = range(0, 100, 10)
    count_personnels = [ages_personnels.count(age) for age in tranches_age]
    count_etudiants = [ages_etudiants.count(age) for age in tranches_age]
    count_enseignants = [ages_enseignants.count(age) for age in tranches_age]




    context = {
        'labels': labels,
        'data': data,
        'male_count': male_count,
        'female_count': female_count,
        'male_countens': male_countens,
        'female_countens': female_countens,
        'missions_encours': missions_encours,
        'autper_encours':autper_encours,
        'autens_encours':autens_encours,
        'autaud_encours':autaud_encours,
        'activeauditeur' : activeauditeur,
        'activeens' : activeens,
        'annee': annee_scolaire_active,
        'tranches_age': tranches_age,
        'count_personnels': count_personnels,
        'count_etudiants': count_etudiants,
        'count_enseignants': count_enseignants,
        
    }

    return render(request, 'Pages/dashboardrh.html', context)

# liste etudiant
@login_required(login_url='/accounts/')
def etu(request):
    etud = Etudiants.objects.all()
    # Récupérer l'année scolaire active
    statuta = Anneescolaires.objects.get(Statutannee='1')
    return render(request,'Pages/etudiants/liste.html', context={'etu':etud,'annee': statuta})

# liste etudiant
@login_required(login_url='/accounts/')
def etureins(request):
    etud = Etudiants.objects.all()
    # Récupérer l'année scolaire active
    statuta = Anneescolaires.objects.get(Statutannee='1')
    return render(request,'Pages/etudiants/listereins.html', context={'etu':etud,'annee': statuta})

# liste etudiant rh
@login_required()
def eturh(request):
    etud = Etudiants.objects.filter(User__statut='ACTIVE')
    # Récupérer l'année scolaire active
    statuta = Anneescolaires.objects.get(Statutannee='1')
    return render(request,'Pages/etudiants/listerh.html', context={'etu':etud,'annee': statuta})




# liste par classe etudiant
@login_required(login_url='/accounts/')
def etuclasse(request, id, iid, iiid):
    etud = Etudiants.objects.filter(Salle=id,Ecole=iid).order_by('Nom','Prenoms')
    # Récupérer l'année scolaire active
    statuta = Anneescolaires.objects.get(Statutannee='1')
    return render(request,'Pages/etudiants/listeclasse.html', context={'etu':etud,'clas':id,'ecol':iid,'ni':iiid,'annee': statuta})

# liste par classe etudiant template
@login_required(login_url='/accounts/')
def etuclassetemplate(request, id, iid):
    etud = Etudiants.objects.filter(Salle=id,Ecole=iid).order_by('Nom','Prenoms')
    nombre_salle = Etudiants.objects.filter(Salle=id,Ecole=iid).count()

    # Récupérer l'année scolaire active
    statuta = Anneescolaires.objects.get(Statutannee='1')
    return render(request,'Pages/etudiants/listeclassetemplate.html', context={'etu':etud,'annee': statuta,'nombre_salle':nombre_salle})

# liste 
@login_required()
def auditeuruser(request):
    auditeur = Utilisateur.objects.filter(role='ETUDIANT')
    # Récupérer l'année scolaire active
    statuta = Anneescolaires.objects.get(Statutannee='1')
    return render(request,'Pages/etudiants/listeutilisateur.html', context={'enseignantliste':auditeur,'annee': statuta})

# liste 
@login_required()
def auditeuruserrh(request):
    auditeur = Utilisateur.objects.filter(role='ETUDIANT')
    # Récupérer l'année scolaire active
    statuta = Anneescolaires.objects.get(Statutannee='1')
    return render(request,'Pages/etudiants/listeutilisateurrh.html', context={'enseignantliste':auditeur,'annee': statuta})

# ajouter etudiant
@login_required()
def add_etudiant(request):
     # Récupérer l'année scolaire active
    statuta = Anneescolaires.objects.get(Statutannee='1')
    if request.method == 'POST':
        fm = EtudiantRegistration(request.POST, request.FILES)
        if fm.is_valid():
            messages.success(request, 'Votre formulaire a été envoyé avec succès !')
            photo = fm.cleaned_data['Photo']
            fm.save()
            return redirect('liste')
    else :
        fm = EtudiantRegistration()
    return render(request,'Pages/etudiants/add.html',{'form':fm,'annee': statuta}  )

# ajouter etudiant
@login_required()
def addd_etudiant(request):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    if request.method == 'POST':
        fm = EtudiantRegistration(request.POST, request.FILES)
        if fm.is_valid():
            messages.success(request, 'Votre formulaire a été envoyé avec succès !')
            photo = fm.cleaned_data['Photo']
           
            student = fm.save(commit=False)
            student.save()
            # Créer un utilisateur avec le même nom d'utilisateur et mot de passe que l'étudiant
            userna = fm.cleaned_data['Nom']
            usernam = fm.cleaned_data['Prenoms']
            FString = userna + " " + usernam
            username = FString
            role = 'ETUDIANT'
            password = 'P@ssword'
            statu='ACTIVE'
            profile_photo=photo
            encoded_password = make_password(password)
            email = fm.cleaned_data['Email']
            userm = Utilisateur.objects.create(statut=statu,profile_photo=profile_photo,username=username,last_name=usernam,first_name=userna, password=encoded_password, email=email,role=role)
            dossier = DossierScolaires.objects.create(Anneescolaire=statuta,Etudiant=student,HeuresabsenceS1=0,HeuresabsenceS2=0,Matricule=student.Matricule,
                Nom=student.Nom,
                Prenoms=student.Prenoms,
                Datenaissance=student.Datenaissance,
                Lieunaissance=student.Lieunaissance,
                Telephone=student.Telephone,
                Contact=student.Contact,
                Email=student.Email,
                Photo=student.Photo,
                Ecole=student.Ecole,
                Sexe=student.Sexe,
                Matrimoniale=student.Matrimoniale,
                Niveau=student.Niveau,
                Salle=student.Salle)
            ucontact = Contacturgences.objects.create(Qualite='Père',Nom=userna,Prenoms=usernam,Etudiant=student,Telephone=userna,Contact="Néant",Email="Néant")
            # Ajouter l'utilisateur à l'objet Etudiant et enregistrer à nouveau
            student.User = userm
            student.save()
            return redirect('liste')
    else :
        matricule = str(uuid.uuid4())[:8] # génération d'une chaîne aléatoire de 8 caractères
        fm = EtudiantRegistration(initial={'Matricule': matricule})
    return render(request,'Pages/etudiants/add.html',{'form':fm,'annee': statuta}  )

@login_required()
def addd_etudiantclasse(request, ecole, clas, niv):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    
    if request.method == 'POST':
        fm = EtudiantRegistration(request.POST, request.FILES)
        if fm.is_valid():
            messages.success(request, 'Votre formulaire a été envoyé avec succès !')
            photo = fm.cleaned_data['Photo']
            nom = fm.cleaned_data['Nom']
            prenom = fm.cleaned_data['Prenoms']
            email = fm.cleaned_data['Email']
            student = fm.save(commit=False)
            student.save()

            
            # Créer un utilisateur avec le même nom d'utilisateur et mot de passe que l'étudiant
            FString = nom + " " + prenom
            username = FString
            role = 'ETUDIANT'
            password = 'P@ssword'
            statu = 'ACTIVE'
            profile_photo = photo
            encoded_password = make_password(password)
            
            userm = Utilisateur.objects.create(statut=statu,
                profile_photo=profile_photo, username=username, last_name=prenom, 
                first_name=nom, password=encoded_password, email=email, role=role
            )
            
            dossier = DossierScolaires.objects.create(
                Anneescolaire=statuta, Etudiant=student, HeuresabsenceS1=0, HeuresabsenceS2=0,Matricule=student.Matricule,
                Nom=student.Nom,
                Prenoms=student.Prenoms,
                Datenaissance=student.Datenaissance,
                Lieunaissance=student.Lieunaissance,
                Telephone=student.Telephone,
                Contact=student.Contact,
                Email=student.Email,
                Photo=student.Photo,
                Ecole=student.Ecole,
                Sexe=student.Sexe,
                Matrimoniale=student.Matrimoniale,
                Niveau=student.Niveau,
                Salle=student.Salle
            )
            ucontact = Contacturgences.objects.create(Qualite='Père',Nom=nom,Prenoms=prenom,Etudiant=student,Telephone="Néant",Contact="Néant",Email="Néant")
            # Ajouter l'utilisateur à l'objet Etudiant et enregistrer à nouveau
            student.User = userm
            student.save()
            
            return redirect(reverse('etuclasse', args=[clas,ecole, niv]))
    else:
        matricule = str(uuid.uuid4())[:8] # génération d'une chaîne aléatoire de 8 caractères
        fm = EtudiantclasseRegistration(initial={'Matricule': matricule, 'Salle': clas, 'Ecole': ecole, 'Niveau': niv,'annee': statuta})
    context = {
        'form': fm,
        'clas': clas,
        'ecole': ecole,
        'niv': niv,
    }
    return render(request, 'Pages/etudiants/addclasse.html', context)


def get_salles(request):
    niveau_id = request.GET.get('niveau_id')
    salles = Salles.objects.filter(Niveau_id=niveau_id).values('id', 'Salle')
    return JsonResponse(list(salles), safe=False)


# modifier dossier auditeur
@login_required()
def auditeurmodif(request, id):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    if request.method == 'POST':
        pi = Etudiants.objects.get(pk=id)
        fm = EtudiantmRegistration(request.POST, request.FILES, instance=pi)
        if fm.is_valid():
            fm.save()
            return redirect('ecoleliste')
    else:
        pi = Etudiants.objects.get(pk=id)
        fm = EtudiantmRegistration(instance=pi)
    return render(request, 'Pages/etudiants/modif.html',{'form':fm,'annee': statuta})

# modifier dossier auditeur
@login_required()
def auditeurmodifreins(request, id):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    if request.META.get('HTTP_X_REQUESTED_WITH') == 'XMLHttpRequest':
        niveau_id = request.GET.get('niveau_id')

        salles = Salles.objects.filter(Niveau__id=niveau_id)

        salle_list = []
        for salle in salles:
            salle_list.append({'Salle': salle.Salle})

        return JsonResponse(salle_list, safe=False)
    
    if request.method == 'POST':
        pi = Etudiants.objects.get(pk=id)
        fm = EtudiantmRegistration(request.POST, request.FILES, instance=pi)
        if fm.is_valid():
            student = fm.save(commit=False)
            dossier_exists = DossierScolaires.objects.filter(Anneescolaire=statuta, Etudiant=student).exists()
            if not dossier_exists:
                dossier = DossierScolaires.objects.create(
                    Anneescolaire=statuta,
                    Etudiant=student,
                    HeuresabsenceS1=0,
                    HeuresabsenceS2=0,
                    Matricule=student.Matricule,
                    Nom=student.Nom,
                    Prenoms=student.Prenoms,
                    Datenaissance=student.Datenaissance,
                    Lieunaissance=student.Lieunaissance,
                    Telephone=student.Telephone,
                    Contact=student.Contact,
                    Email=student.Email,
                    Photo=student.Photo,
                    Ecole=student.Ecole,
                    Sexe=student.Sexe,
                    Matrimoniale=student.Matrimoniale,
                    Niveau=student.Niveau,
                    Salle=student.Salle
                )
                messages.success(request, "Le dossier scolaire a été créé avec succès.")
            else:
                messages.error(request, "Le dossier scolaire existe déjà pour cet étudiant dans la même année scolaire.")
            fm.save()
            return redirect('ecoleliste')
    else:
        pi = Etudiants.objects.get(pk=id)
        fm = EtudiantmRegistration(instance=pi)
    return render(request, 'Pages/etudiants/modifreins1.html', {'form': fm,'annee': statuta})


# modifier dossier auditeur total
@login_required()
def auditeurmodiftotal(request, id):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    if request.method == 'POST':
        pi = Etudiants.objects.get(pk=id)
        fm = EtudiantRegistration(request.POST, instance=pi)
        if fm.is_valid():
            fm.save()
            return redirect('ecoleliste')
    else:
        pi = Etudiants.objects.get(pk=id)
        fm = EtudiantRegistration(instance=pi)
    return render(request, 'Pages/etudiants/modiftotal.html',{'form':fm,'annee': statuta})

# ECOLE........................................................................................................................................
# liste 
@login_required()
def ecoleliste(request):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    etablissement = Ecoles.objects.all()
    return render(request,'Pages/ecoles/liste.html', context={'ecoleliste':etablissement,'annee': statuta})

@login_required()
def ecolelistecat(request):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    etablissement = Ecoles.objects.all()
    return render(request,'Pages/ecoles/listecat.html', context={'ecoleliste':etablissement,'annee': statuta})

# ajouter
@login_required()
def ecoleajout(request):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    if request.method == 'POST':
        fm = EcoleRegistration(request.POST, request.FILES)
        if fm.is_valid():
            image = fm.cleaned_data['image']
            fm.save()
            messages.success(request, 'Votre formulaire a été envoyé avec succès !')
            fm = EcoleRegistration()
            return redirect('ecoleliste')
    else :
        fm = EcoleRegistration()
    return render(request,'Pages/ecoles/ajout.html',{'form':fm,'annee': statuta}  )

# modifier
@login_required()
def ecolemodif(request, id):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    if request.method == 'POST':
        pi = Ecoles.objects.get(pk=id)
        fm = EcoleRegistration(request.POST, instance=pi)
        if fm.is_valid():
            fm.save()
            return redirect('ecoleliste')
    else:
        pi = Ecoles.objects.get(pk=id)
        fm = EcoleRegistration(instance=pi)
    return render(request, 'Pages/ecoles/modif.html',{'form':fm,'annee': statuta})

# supprimer
@login_required()
def delete_object(request, id):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    pi = Ecoles.objects.get(pk=id)
    if request.method == 'POST':
        pi.delete()
        return redirect('ecoleliste')  # rediriger vers la liste des objets

    # Afficher le formulaire de confirmation de suppression
    context = {'object': pi,'annee': statuta}
    return render(request, 'myapp/delete_object.html', context)


# SALLE DE CLASSE------------------------------------------------------------------------------------------------------------
# ajouter classe
@login_required()
def add_classe(request,id):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    salle = Ecoles.objects.get(id=id)
    if request.method == 'POST':
        fm = SalleRegistration(request.POST, request.FILES)
        if fm.is_valid():
            messages.success(request, 'Votre formulaire a été envoyé avec succès !')
            Sal = fm.cleaned_data['Salle']
            Nbre = fm.cleaned_data['Nbreplace']
            Prof = fm.cleaned_data['Professeurprincipal']
            Niv = fm.cleaned_data['Niveau']
            Planning = fm.cleaned_data['Planning']
            grade = Salles.objects.create( Nbreplace=Nbre,Planning=Planning,Niveau=Niv, Anneescolaire=statuta, Salle=Sal, Etablissement=salle, Professeurprincipal=Prof)
            return redirect('ecoleliste')
    else :
        fm = SalleRegistration()
    return render(request,'Pages/salles/add.html',{'form':fm,'annee': statuta}  )

# modifier
@login_required()
def classmodif(request, id):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    if request.method == 'POST':
        pi = Salles.objects.get(pk=id)
        fm = SallemRegistration(request.POST, instance=pi)
        if fm.is_valid():
            fm.save()
            return redirect('ecoleliste')
    else:
        pi = Salles.objects.get(pk=id)
        fm = SallemRegistration(instance=pi)
    return render(request, 'Pages/salles/modif.html',{'form':fm,'annee': statuta})

# Effectif Classe
@login_required() 
def listesalle(request, object_id):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    # Récupérer l'objet à partir de l'ID fourni
    salle = Salles.objects.get(id=object_id)
    
    # Récupérer les données triées en fonction de l'objet
    data = Salles.objects.filter(object=salle).order_by('Niveau')

    # Passer les données triées en contexte au template
    context = {'object': salle, 'data': data,'annee': statuta}
    return render(request, 'Pages/salles/salleecole.html', context)

def object_table(request, id):
   
    statuta = Anneescolaires.objects.get(Statutannee='1')
    # Récupérer les données triées en fonction de l'objet
    queryset = Salles.objects.filter(Etablissement=id, Anneescolaire=statuta).order_by('Niveau')

    # Passer les données triées en contexte au template
    context = {'queryset': queryset,'annee': statuta}
    return render(request, 'Pages/salles/salleecole.html', context)



# DOCUMENT DE CLASSE------------------------------------------------------------------------------------------------------------
# Liste Document 
@login_required()
def listedoc(request):
   
    statuta = Anneescolaires.objects.get(Statutannee='1')
    # Récupérer les données triées en fonction de l'objet
    queryse = Documents.objects.all()

    # Passer les données triées en contexte au template
    context = {'queryse': queryse,'annee': statuta}
    return render(request, 'Pages/documents/listedoc.html', context)

# Ajouter Document
@login_required()
def add_document(request):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    
    if request.method == 'POST':
        
        form = DocumentForm(request.POST, request.FILES)
        if form.is_valid():
            Fil = form.cleaned_data['Cours']
            Typ = Matieres.objects.get(Matiere=Fil)
            File = form.cleaned_data['File']
            Title = form.cleaned_data['Title']
            ids = form.cleaned_data['Salle']
            student = form.save(commit=False)
            userm = Documents.objects.create(Anneescolaire=statuta,Cours=Fil,File=File,Title=Title,Salle=ids,Typematiere=Typ.Typematiere)
            userm.save()
            
            return redirect('documentlist')
    else:
        
        form = DocumentForm()
    return render(request, 'Pages/documents/ajoutdoc.html', {'form': form,'annee': statuta})


# Detail
@login_required()
def document_detail(request, document_id):
    document = Documents.objects.get(id=document_id)
    return FileResponse(document.File)

# Telecharger
@login_required()
def document_download(request, document_id):
    document = Documents.objects.get(id=document_id)
    response = HttpResponse(document.File, content_type='application/force-download')
    response['Content-Disposition'] = f'attachment; filename="{document.File.name}"'
    return response


# ANNEES SCOLAIRES------------------------------------------------------------------------------------------------------------
# liste 
@login_required()
def anneescolaireliste(request):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    anneescolaire = Anneescolaires.objects.all()
    return render(request,'Pages/anneescolaires/liste.html', context={'anneescolaireliste':anneescolaire,'annee': statuta})

# ajouter
@login_required()
def anneescolaireajout(request):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    if request.method == 'POST':
        fm = AnneescolaireRegistration(request.POST)
        if fm.is_valid():
            
            fm.save()
            messages.success(request, 'Votre formulaire a été envoyé avec succès !')
            fm = AnneescolaireRegistration()
            return redirect('anneescolaireliste')
    else :
        fm = AnneescolaireRegistration()
    return render(request,'Pages/anneescolaires/ajout.html',{'form':fm,'annee': statuta}  )

# modifier
@login_required()
def annemodif(request, id):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    if request.method == 'POST':
        pi = Anneescolaires.objects.get(pk=id)
        fm = AnneescolaireRegistration(request.POST, instance=pi)
        if fm.is_valid():
            fm.save()
            return redirect('anneescolaireliste')
    else:
        pi = Anneescolaires.objects.get(pk=id)
        fm = AnneescolaireRegistration(instance=pi)
    return render(request, 'Pages/anneescolaires/modif.html',{'form':fm,'annee': statuta})

# supprimer
@login_required()
def delete_annee(request, id):
    annee = Anneescolaires.objects.get(pk=id)
    annee.delete()
    return redirect('anneescolaireliste')

# Modifier statut
@login_required()
def anneescolairestatut(request, id):
    statut_actif = Statutannees.objects.get(Statutannee='ACTIVE')
    statut_inactif = Statutannees.objects.get(Statutannee='NON ACTIVE')
    annee_statut = Anneescolaires.objects.get(id=id)
    if annee_statut.Statutannee == statut_actif:
        annee_statut.Statutannee = statut_inactif
    else:
        annee_statut.Statutannee = statut_actif
    annee_statut.save()
    return redirect('anneescolaireliste')


# ENSEIGNANTS------------------------------------------------------------------------------------------------------------
# liste 
@login_required()
def enseignantliste(request):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    enseignant = Enseignants.objects.all()
    return render(request,'Pages/enseignants/liste.html', context={'enseignantliste':enseignant,'annee': statuta})

@login_required()
def enseignantlistepers(request):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    enseignant = Enseignants.objects.filter(User__statut='ACTIVE')
    return render(request,'Pages/enseignants/listepers.html', context={'enseignantliste':enseignant,'annee': statuta})

# liste 
@login_required()
def enseignantuser(request):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    enseignant = Utilisateur.objects.filter(role='ENSEIGNANT')
    return render(request,'Pages/enseignants/listeutilisateur.html', context={'enseignantliste':enseignant,'annee': statuta})

@login_required()
def enseignantuserrh(request):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    enseignant = Utilisateur.objects.filter(role='ENSEIGNANT')
    return render(request,'Pages/enseignants/listeutilisateurrh.html', context={'enseignantliste':enseignant,'annee': statuta})

# ajouter
@login_required()
def enseignantajout(request):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    if request.method == 'POST':
        fm = EnseignantRegistration(request.POST, request.FILES)
        if fm.is_valid():
            Photo = fm.cleaned_data['Photo']
            student = fm.save(commit=False)
            student.save()
            # Créer un utilisateur avec le même nom d'utilisateur et mot de passe que l'étudiant
            userna = fm.cleaned_data['Nom']
            usernam = fm.cleaned_data['Prenoms']
            FString = userna + " " + usernam
            username = FString
            role = 'ENSEIGNANT'
            statu = 'ACTIVE'
            password = 'P@ssword'
            profile_photo = Photo
            encoded_password = make_password(password)
            email = fm.cleaned_data['Email']
            userm = Utilisateur.objects.create(statut=statu,profile_photo=profile_photo,username=username,last_name=usernam,first_name=userna, password=encoded_password, email=email,role=role)

            # Ajouter l'utilisateur à l'objet Etudiant et enregistrer à nouveau
            student.User = userm
            student.save()
            return redirect('enseignantliste')
    else :
        matricule = str(uuid.uuid4())[:8] # génération d'une chaîne aléatoire de 8 caractères
        fm = EnseignantRegistration(initial={'Matricule': matricule})
    return render(request,'Pages/enseignants/ajout.html',{'form':fm,'annee': statuta}  )

# modifier
def enseignantmodif(request, id):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    if request.method == 'POST':
        pi = Enseignants.objects.get(pk=id)
        fm = EnseignantRegistration(request.POST, instance=pi)
        if fm.is_valid():
            fm.save()
            return redirect('enseignantliste')
    else:
        pi = Enseignants.objects.get(pk=id)
        fm = EnseignantRegistration(instance=pi)
    return render(request, 'Pages/enseignants/modif.html',{'form':fm,'annee': statuta})

# supprimer
@login_required()
def delete_ensignant(request, id):
    enseignant = Enseignants.objects.get(pk=id)
    enseignant.delete()
    return redirect('enseignantliste')

# NOTES------------------------------------------------------------------------------------------------------------
@login_required() # type: ignore
def add_grade(request,object_id,objec_id):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    if request.method == 'POST':
        statuta = Anneescolaires.objects.get(Statutannee='1')
        statuti = Decoupages.objects.get(Statutdecoup='1')
        course_id = request.POST['course']
        score = request.POST.get('score', None)
        if score is not None:
            score = float(score)
            if 19.0 <= score <= 20.0:
                
                
                app = "Excellent"
                student = Etudiants.objects.get(id=object_id)
                course = Matieres.objects.get(id=course_id)

                if Notes.objects.filter(Etudiant=student, Cours=course, Anneescolaire=statuta, Decoupage=statuti).exists():
                    # Affichage d'un message d'erreur
                    messages.warning(request, "Cet étudiant existe déjà dans la base de données.")
                    return redirect(reverse('notebase', args=[object_id,objec_id]))
                   

                grade = Notes.objects.create( Etudiant=student,Appreciation=app, Cours=course, Note=score, Anneescolaire=statuta, Decoupage=statuti, Salle=student.Salle,  Professeur=course.Professeur, Typematiere=course.Typematiere)
                messages.success(request, f"Grade added for {course.Matiere}")
                return redirect('ecoleliste')
            elif 17.0 <= score <= 18.75:
                app = "Très Bien"
                
                student = Etudiants.objects.get(id=object_id)
                course = Matieres.objects.get(id=course_id)

                if Notes.objects.filter(Etudiant=student, Cours=course, Anneescolaire=statuta, Decoupage=statuti).exists():
                    # Affichage d'un message d'erreur
                    messages.warning(request, "Cet étudiant existe déjà dans la base de données.")
                    return redirect(reverse('notebase', args=[object_id,objec_id]))

                grade = Notes.objects.create( Etudiant=student,Appreciation=app, Cours=course, Note=score, Anneescolaire=statuta, Decoupage=statuti, Salle=student.Salle,  Professeur=course.Professeur, Typematiere=course.Typematiere)
                messages.success(request, f"Grade added for {course.Matiere}")
                return redirect('ecoleliste')
            elif 15.0 <= score <= 16.99:
                app = "Bien"
                
                student = Etudiants.objects.get(id=object_id)
                course = Matieres.objects.get(id=course_id)

                if Notes.objects.filter(Etudiant=student, Cours=course, Anneescolaire=statuta, Decoupage=statuti).exists():
                    # Affichage d'un message d'erreur
                    messages.warning(request, "Cet étudiant existe déjà dans la base de données.")
                    return redirect(reverse('notebase', args=[object_id,objec_id]))
                
                grade = Notes.objects.create( Etudiant=student,Appreciation=app, Cours=course, Note=score, Anneescolaire=statuta, Decoupage=statuti, Salle=student.Salle,  Professeur=course.Professeur, Typematiere=course.Typematiere)
                messages.success(request, f"Grade added for {course.Matiere}")
                return redirect('ecoleliste')
            elif 13.0 <= score <= 14.99:
                app = "Assez Bien"
                
                student = Etudiants.objects.get(id=object_id)
                course = Matieres.objects.get(id=course_id)

                if Notes.objects.filter(Etudiant=student, Cours=course, Anneescolaire=statuta, Decoupage=statuti).exists():
                    # Affichage d'un message d'erreur
                    messages.warning(request, "Cet étudiant existe déjà dans la base de données.")
                    return redirect(reverse('notebase', args=[object_id,objec_id]))
                
                grade = Notes.objects.create( Etudiant=student,Appreciation=app, Cours=course, Note=score, Anneescolaire=statuta, Decoupage=statuti, Salle=student.Salle,  Professeur=course.Professeur, Typematiere=course.Typematiere)
                messages.success(request, f"Grade added for {course.Matiere}")
                return redirect('ecoleliste')
            elif 11.0 <= score <= 12.99:
                app = "Passable"
                
                student = Etudiants.objects.get(id=object_id)
                course = Matieres.objects.get(id=course_id)

                if Notes.objects.filter(Etudiant=student, Cours=course, Anneescolaire=statuta, Decoupage=statuti).exists():
                    # Affichage d'un message d'erreur
                    messages.warning(request, "Cet étudiant existe déjà dans la base de données.")
                    return redirect(reverse('notebase', args=[object_id,objec_id]))
                
                grade = Notes.objects.create( Etudiant=student,Appreciation=app, Cours=course, Note=score, Anneescolaire=statuta, Decoupage=statuti, Salle=student.Salle,  Professeur=course.Professeur, Typematiere=course.Typematiere)
                messages.success(request, f"Grade added for {course.Matiere}")
                return redirect('ecoleliste')
            elif 0 <= score <= 10.99:
                app = "Faible"
               
                student = Etudiants.objects.get(id=object_id)
                course = Matieres.objects.get(id=course_id)

                if Notes.objects.filter(Etudiant=student, Cours=course, Anneescolaire=statuta, Decoupage=statuti).exists():
                    # Affichage d'un message d'erreur
                    messages.warning(request, "Cet étudiant existe déjà dans la base de données.")
                    return redirect(reverse('notebase', args=[object_id,objec_id]))
                
                grade = Notes.objects.create( Etudiant=student,Appreciation=app, Cours=course, Note=score, Anneescolaire=statuta, Decoupage=statuti, Salle=student.Salle,  Professeur=course.Professeur, Typematiere=course.Typematiere)
                messages.success(request, f"Grade added for {course.Matiere}")
                return redirect('ecoleliste')
            else:
                app = "Score out of range"
        else:
            app = "No score provided"
        return app

        
    else:
        courses = Matieres.objects.filter(Salle=objec_id).order_by('Typematiere')
        context = {
            'courses': courses,'annee': statuta
        }
        return render(request, 'Pages/bulletins/ajoutnote.html', context)
    
# ajouter moyenne plusieurs

def add_note(request, object_id, objec_id):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    if request.method == 'POST':
        statut = Anneescolaires.objects.get(Statutannee='1')
        decoupage = Decoupages.objects.get(Statutdecoup='1')
        courses_ids = request.POST.getlist('course')
        score = request.POST.get('score', None)
        if score is not None:
            score = float(score)
            if 19.0 <= score <= 20.0:
                app = "Excellent"
            elif 17.0 <= score <= 18.75:
                app = "Très Bien"
            elif 15.0 <= score <= 16.99:
                app = "Bien"
            elif 13.0 <= score <= 14.99:
                app = "Assez Bien"
            elif 11.0 <= score <= 12.99:
                app = "Passable"
            else:
                app = "Insuffisant"
                
            student = Etudiants.objects.get(id=object_id)
            courses = Matieres.objects.filter(id__in=courses_ids)
            
            for course in courses:
                if Notes.objects.filter(Etudiant=student, Cours=course, Anneescolaire=statut, Decoupage=decoupage).exists():
                    # Affichage d'un message d'erreur
                    messages.warning(request, "Cet étudiant a déjà une note pour ce cours.")
                else:
                    grade = Notes.objects.create(Etudiant=student, Appreciation=app, Cours=course, Note=score, Anneescolaire=statut, Decoupage=decoupage, Salle=student.Salle, Professeur=course.Professeur, Typematiere=course.Typematiere)
                    messages.success(request, f"Grade added for {course.Matiere}")
            
            return redirect('ecoleliste')
    else:
        courses = Matieres.objects.filter(Salle=object_id).order_by('Typematiere')
        context = {
            'courses': courses,'annee': statuta
        }
        return render(request, 'Pages/bulletins/addnote.html', context)



# ajouter moyenne auditeur espace enseignant
@login_required() # type: ignore
def add_grade2(request,object_id,objec_id):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    if request.method == 'POST':
        statuta = Anneescolaires.objects.get(Statutannee='1')
        statuti = Decoupages.objects.get(Statutdecoup='1')
        course_id = request.POST['course']
        score = request.POST.get('score', None)
        if score is not None:
            score = float(score)
            if 19.0 <= score <= 20.0:
                
                
                app = "Excellent"
                student = Etudiants.objects.get(id=object_id)
                course = Matieres.objects.get(id=course_id)

                if Notes.objects.filter(Etudiant=student, Cours=course, Anneescolaire=statuta, Decoupage=statuti).exists():
                    # Affichage d'un message d'erreur
                    messages.warning(request, "Cet étudiant existe déjà dans la base de données.")
                    return redirect(reverse('notebase2', args=[object_id,objec_id]))
                   

                grade = Notes.objects.create( Etudiant=student,Appreciation=app, Cours=course, Note=score, Anneescolaire=statuta, Decoupage=statuti, Salle=student.Salle,  Professeur=course.Professeur, Typematiere=course.Typematiere)
                messages.success(request, f"Grade added for {course.Matiere}")
                return redirect('ecoleliste')
            elif 17.0 <= score <= 18.75:
                app = "Très Bien"
                
                student = Etudiants.objects.get(id=object_id)
                course = Matieres.objects.get(id=course_id)

                if Notes.objects.filter(Etudiant=student, Cours=course, Anneescolaire=statuta, Decoupage=statuti).exists():
                    # Affichage d'un message d'erreur
                    messages.warning(request, "Cet étudiant existe déjà dans la base de données.")
                    return redirect(reverse('notebase2', args=[object_id,objec_id]))

                grade = Notes.objects.create( Etudiant=student,Appreciation=app, Cours=course, Note=score, Anneescolaire=statuta, Decoupage=statuti, Salle=student.Salle,  Professeur=course.Professeur, Typematiere=course.Typematiere)
                messages.success(request, f"Grade added for {course.Matiere}")
                return redirect('ecoleliste')
            elif 15.0 <= score <= 16.99:
                app = "Bien"
                
                student = Etudiants.objects.get(id=object_id)
                course = Matieres.objects.get(id=course_id)

                if Notes.objects.filter(Etudiant=student, Cours=course, Anneescolaire=statuta, Decoupage=statuti).exists():
                    # Affichage d'un message d'erreur
                    messages.warning(request, "Cet étudiant existe déjà dans la base de données.")
                    return redirect(reverse('notebase2', args=[object_id,objec_id]))
                
                grade = Notes.objects.create( Etudiant=student,Appreciation=app, Cours=course, Note=score, Anneescolaire=statuta, Decoupage=statuti, Salle=student.Salle,  Professeur=course.Professeur, Typematiere=course.Typematiere)
                messages.success(request, f"Grade added for {course.Matiere}")
                return redirect('ecoleliste')
            elif 13.0 <= score <= 14.99:
                app = "Assez Bien"
                
                student = Etudiants.objects.get(id=object_id)
                course = Matieres.objects.get(id=course_id)

                if Notes.objects.filter(Etudiant=student, Cours=course, Anneescolaire=statuta, Decoupage=statuti).exists():
                    # Affichage d'un message d'erreur
                    messages.warning(request, "Cet étudiant existe déjà dans la base de données.")
                    return redirect(reverse('notebase2', args=[object_id,objec_id]))
                
                grade = Notes.objects.create( Etudiant=student,Appreciation=app, Cours=course, Note=score, Anneescolaire=statuta, Decoupage=statuti, Salle=student.Salle,  Professeur=course.Professeur, Typematiere=course.Typematiere)
                messages.success(request, f"Grade added for {course.Matiere}")
                return redirect('ecoleliste')
            elif 11.0 <= score <= 12.99:
                app = "Passable"
                
                student = Etudiants.objects.get(id=object_id)
                course = Matieres.objects.get(id=course_id)

                if Notes.objects.filter(Etudiant=student, Cours=course, Anneescolaire=statuta, Decoupage=statuti).exists():
                    # Affichage d'un message d'erreur
                    messages.warning(request, "Cet étudiant existe déjà dans la base de données.")
                    return redirect(reverse('notebase2', args=[object_id,objec_id]))
                
                grade = Notes.objects.create( Etudiant=student,Appreciation=app, Cours=course, Note=score, Anneescolaire=statuta, Decoupage=statuti, Salle=student.Salle,  Professeur=course.Professeur, Typematiere=course.Typematiere)
                messages.success(request, f"Grade added for {course.Matiere}")
                return redirect('ecoleliste')
            elif 0 <= score <= 10.99:
                app = "Faible"
               
                student = Etudiants.objects.get(id=object_id)
                course = Matieres.objects.get(id=course_id)

                if Notes.objects.filter(Etudiant=student, Cours=course, Anneescolaire=statuta, Decoupage=statuti).exists():
                    # Affichage d'un message d'erreur
                    messages.warning(request, "Cet étudiant existe déjà dans la base de données.")
                    return redirect(reverse('notebase2', args=[object_id,objec_id]))
                
                grade = Notes.objects.create( Etudiant=student,Appreciation=app, Cours=course, Note=score, Anneescolaire=statuta, Decoupage=statuti, Salle=student.Salle,  Professeur=course.Professeur, Typematiere=course.Typematiere)
                messages.success(request, f"Grade added for {course.Matiere}")
                return redirect('ecoleliste')
            else:
                app = "Score out of range"
        else:
            app = "No score provided"
        return app

        
    else:
        courses = Matieres.objects.filter(Salle=objec_id).order_by('Typematiere')
        context = {
            'courses': courses,'annee': statuta
        }
        return render(request, 'Pages/bulletins/ajoutnoteespace.html', context)
    


# modifier moyenne auditeur
@login_required() # type: ignore
def moymodiff(request, id):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    if request.method == 'POST':
        pi = Notes.objects.get(pk=id)
        fm = MoyennemRegistration(request.POST, instance=pi)
        if fm.is_valid():
            score = fm.cleaned_data['Note']
            if score is not None:
                score = float(score)
                if 19.0 <= score <= 20.0:
                    audmoy = fm.save(commit=False)
                    app = "Excellent"
                    audmoy.Appreciation=app
                    audmoy.save()
                    return redirect('ecoleliste')
                elif 17.0 <= score <= 18.75:
                    audmoy = fm.save(commit=False)
                    app = "Très Bien"
                    audmoy.Appreciation=app
                    audmoy.save()
                    return redirect('ecoleliste')
                elif 15.0 <= score <= 16.99:
                    audmoy = fm.save(commit=False)
                    app = "Bien"
                    audmoy.Appreciation=app
                    audmoy.save()
                    return redirect('ecoleliste')
                elif 13.0 <= score <= 14.99:
                    audmoy = fm.save(commit=False)
                    app = "Assez Bien"
                    audmoy.Appreciation=app
                    audmoy.save()
                    return redirect('ecoleliste')
                elif 11.0 <= score <= 12.99:
                    audmoy = fm.save(commit=False)
                    app = "Passable"
                    audmoy.Appreciation=app
                    audmoy.save()
                    return redirect('ecoleliste')
                elif 0 <= score <= 10.99:
                    audmoy = fm.save(commit=False)
                    app = "Faible"
                    audmoy.Appreciation=app
                    audmoy.save()
                    return redirect('ecoleliste')
                else:
                    app = "Score out of range"
            else:
                app = "No score provided"
            return app
            
    else:
        pi = Notes.objects.get(pk=id)
        fm = MoyennemRegistration(instance=pi)
    return render(request, 'Pages/bulletins/modif.html',{'form':fm,'annee': statuta})


# Afficher liste Moyenne par auditeurs
@login_required()
def home(request,id):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    student = Etudiants.objects.get(id=id)
    statuti = Decoupages.objects.get(Statutdecoup='1')
    statuta = Anneescolaires.objects.get(Statutannee='1')
    grades = Notes.objects.filter(Etudiant=student, Decoupage=statuti, Anneescolaire=statuta)
    context = {
        'student': student,
        'grades': grades,
        'annee': statuta
    }
    return render(request, 'Pages/bulletins/listenote.html', context)

# Afficher liste Moyenne par auditeurs
@login_required()
def home2(request,id):
    student = Etudiants.objects.get(id=id)
    statuti = Decoupages.objects.get(Statutdecoup='1')
    statuta = Anneescolaires.objects.get(Statutannee='1')
    grades = Notes.objects.filter(Etudiant=student, Decoupage=statuti, Anneescolaire=statuta)
    context = {
        'student': student,
        'grades': grades,'annee': statuta
    }
    return render(request, 'Pages/bulletins/listenoteespace.html', context)

# MOYENNES------------------------------------------------------------------------------------------------------------
# Calcul moyennes
@login_required()
def calculate_results(request, id, iid):
    # Obtenez l'année scolaire et la division actives
    statuta = Anneescolaires.objects.get(Statutannee='1')
    statuti = Decoupages.objects.get(Statutdecoup='1')

    # Obtenez tous les étudiants dans la salle donnée appartenant à l'école donnée
    students = Etudiants.objects.filter(Salle=id, Ecole=iid)

    # Boucle sur tous les étudiants
    for student in students:
        # Obtenez les résultats de l'étudiant pour l'année scolaire et la division actives
        results = Notes.objects.filter(Etudiant=student, Anneescolaire=statuta, Decoupage=statuti)

        # Calculer la somme des notes de l'étudiant
        total_score = 0
        for result in results:
            total_score += result.Note # type: ignore

       

        # Calculer la moyenne de l'étudiant
            if results:
                avg_score = total_score / len(results)
            else:
                avg_score = 0

        # Enregistrer la moyenne de l'étudiant dans la base de données
                grade = Moyennes.objects.create(Etudiant=student, Moyenne=avg_score, Total=total_score, Anneescolaire=statuta, Decoupage=statuti)

    # Obtenez tous les étudiants triés par leur note moyenne
    ranked_students = Moyennes.objects.filter(Anneescolaire=statuta, Decoupage=statuti).order_by('-Moyenne')

    # Attribuez un rang à chaque étudiant en fonction de sa note moyenne
    rank = 1
    for student in ranked_students:
        student.Rang = rank # type: ignore
        student.save()
        rank += 1

    return redirect('ecoleliste')

# Afficher les bulletins de la classe
@login_required()
def homebu(request, id):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    statuti = Decoupages.objects.get(Statutdecoup='1')

    books = Moyennes.objects.select_related('Etudiant').all()
    data = []
    for book in books:
        etudiantn = book.Etudiant.Nom
        etudiantp = book.Etudiant.Prenoms
        etudiants = book.Etudiant.Salle
        moyenne = book.Moyenne
        moyenned = book.Decoupage
        moyennea = book.Anneescolaire
        moyennet = book.Total
        moyenner = book.Rang
        data.append({
            'etudiantn': etudiantn,
            'etudiantp': etudiantp,
            'etudiants': etudiants,
            'moyenned': moyenned,
            'moyennea': moyennea,
            'moyennet': moyennet,
            'moyenner': moyenner,
            'moyenne': moyenne
        })

    etudiant = Etudiants.objects.prefetch_related('Salle').get(id=id)

    notes = Notes.objects.filter(Etudiant=etudiant, Anneescolaire=statuta, Decoupage=statuti)
    
    dat = []
    
    for note in notes:
        notem = note.Cours
        notev = note.Note
        notep = note.Professeur
        notea = note.Appreciation
        notet = note.Typematiere

        dat.append({
            'notem': notem,
            'notep': notep,
            'notea': notea,
            'notet': notet,
            'notev': notev
        })

    return render(request, 'Pages/bulletins/resultat.html', context={'data': data, 'dat': dat,'annee': statuta})

    


# DOCUMENTS ESPACE ENSEIGNANT------------------------------------------------------------------------------------------------------------
# Ajouter
@login_required()
def add_documente(request):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    if request.method == 'POST':
        form = DocumentForm(request.POST, request.FILES)
        if form.is_valid():
            form.save()
            return redirect('documentlist')
    else:
        form = DocumentForm()
    return render(request, 'add_document.html', {'form': form,'annee': statuta})

# modifier
@login_required()
def docmodif(request, id):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    if request.method == 'POST':
        pi = Documents.objects.get(pk=id)
        fm = DocumentForm(request.POST, instance=pi)
        if fm.is_valid():
            fm.save()
            return redirect('documentlist')
    else:
        pi = Documents.objects.get(pk=id)
        fm = DocumentForm(instance=pi)
    return render(request, 'Pages/documents/modif.html',{'form':fm,'annee': statuta})

# supprimer
@login_required()
def delete_doc(request, id):
    doc = Documents.objects.get(pk=id)
    doc.delete()
    return redirect('documentlist')


# Detail
@login_required()
def document_detaile(request, document_id):
    document = Documents.objects.get(id=document_id)
    return FileResponse(document.File)

# Ajouter
@login_required()
def document_downloade(request, document_id):
    document = Documents.objects.get(id=document_id)
    response = HttpResponse(document.File, content_type='application/force-download')
    response['Content-Disposition'] = f'attachment; filename="{document.File.name}"'
    return response

# DOCUMENTS ESPACE ENSEIGNANT------------------------------------------------------------------------------------------------------------
# Ajouter



# TYPES INSTITUTIONS------------------------------------------------------------------------------------------------------------
# liste 
@login_required()
def typeinstliste(request):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    typeinst = Typeecoles.objects.all()
    return render(request,'Pages/typeecoles/liste.html', context={'typeinst':typeinst,'annee': statuta})

# ajouter
@login_required()
def typeinsteajout(request):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    if request.method == 'POST':
        
        fm = TypeinstRegistration(request.POST)
        if fm.is_valid():
            
            fm.save()
            messages.success(request, 'Votre formulaire a été envoyé avec succès !')
            fm = TypeinstRegistration()
            return redirect('typeinstliste')
    else :
        fm = TypeinstRegistration()
    return render(request,'Pages/typeecoles/ajout.html',context={'form':fm,'annee': statuta}  )

# modifier
@login_required()
def typeecolemodif(request, id):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    if request.method == 'POST':
        pi = Typeecoles.objects.get(pk=id)
        fm = TypeinstRegistration(request.POST, instance=pi)
        if fm.is_valid():
            fm.save()
            return redirect('typeinstliste')
    else:
        pi = Typeecoles.objects.get(pk=id)
        fm = TypeinstRegistration(instance=pi)
    return render(request, 'Pages/typeecoles/modif.html',context={'form':fm,'annee': statuta})

# supprimer
@login_required()
def delete_typeecole(request, id):
    typeecole = Typeecoles.objects.get(pk=id)
    typeecole.delete()
    return redirect('typeinstliste')

# TYPES MATIÈRES------------------------------------------------------------------------------------------------------------
# liste 
@login_required()
def typematliste(request):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    typemat = Typematieres.objects.all()
    return render(request,'Pages/typematieres/liste.html', context={'typemat':typemat,'annee': statuta})

# ajouter
@login_required()
def typematajout(request):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    if request.method == 'POST':
        fm = TypematRegistration(request.POST)
        if fm.is_valid():
            
            fm.save()
            messages.success(request, 'Votre formulaire a été envoyé avec succès !')
            fm = TypematRegistration()
            return redirect('typematliste')
    else :
        fm = TypematRegistration()
    return render(request,'Pages/typematieres/ajout.html',{'form':fm,'annee': statuta}  )

# modifier
@login_required()
def typematmodif(request, id):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    if request.method == 'POST':
        pi = Typematieres.objects.get(pk=id)
        fm = TypematRegistration(request.POST, instance=pi)
        if fm.is_valid():
            fm.save()
            return redirect('typematliste')
    else:
        pi = Typematieres.objects.get(pk=id)
        fm = TypematRegistration(instance=pi)
    return render(request, 'Pages/typematieres/modif.html',{'form':fm,'annee': statuta})

# supprimer
@login_required()
def delete_typemat(request, id):
    typemat = Typematieres.objects.get(pk=id)
    typemat.delete()
    return redirect('typematliste')

# MATIÈRES------------------------------------------------------------------------------------------------------------
# liste 
@login_required()
def matliste(request):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    mat = Matieres.objects.all()
    return render(request,'Pages/matieres/liste.html', context={'mat':mat,'annee': statuta})

# ajouter
@login_required()
def matajout(request):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    if request.method == 'POST':
        fm = MatiereRegistration(request.POST)
        if fm.is_valid():
            fm.save()
            messages.success(request, 'Votre formulaire a été envoyé avec succès !')
            fm = MatiereRegistration()
            return redirect('matliste')
    else :
        fm = MatiereRegistration()
    return render(request,'Pages/matieres/ajout.html',{'form':fm,'annee': statuta}  )

# modifier
@login_required()
def matmodif(request, id):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    if request.method == 'POST':
        pi = Matieres.objects.get(pk=id)
        fm = MatiereRegistration(request.POST, instance=pi)
        if fm.is_valid():
            fm.save()
            return redirect('matliste')
    else:
        pi = Matieres.objects.get(pk=id)
        fm = MatiereRegistration(instance=pi)
    return render(request, 'Pages/matieres/modif.html',{'form':fm,'annee': statuta})

# supprimer
@login_required()
def delete_mat(request, id):
    mat = Matieres.objects.get(pk=id)
    mat.delete()
    return redirect('matliste')


# DECOUPAGES------------------------------------------------------------------------------------------------------------
# liste 
@login_required()
def decoupliste(request):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    decoup = Decoupages.objects.all()
    return render(request,'Pages/decoupages/liste.html', context={'decoup':decoup,'annee': statuta})

# ajouter
@login_required()
def decoupajout(request):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    if request.method == 'POST':
        fm = SessionRegistration(request.POST)
        if fm.is_valid():
            fm.save()
            messages.success(request, 'Votre formulaire a été envoyé avec succès !')
            fm = SessionRegistration()
            return redirect('decoupliste')
    else :
        fm = SessionRegistration()
    return render(request,'Pages/decoupages/ajout.html',{'form':fm,'annee': statuta}  )

# modifier
@login_required()

def decoupmodif(request, id):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    if request.method == 'POST':
        pi = Decoupages.objects.get(pk=id)
        fm = SessionRegistration(request.POST, instance=pi)
        if fm.is_valid():
            fm.save()
            return redirect('decoupliste')
    else:
        pi = Decoupages.objects.get(pk=id)
        fm = SessionRegistration(instance=pi)
    return render(request, 'Pages/decoupages/modif.html',{'form':fm,'annee': statuta})

# supprimer
@login_required()
def delete_decoup(request, id):
    mat = Decoupages.objects.get(pk=id)
    mat.delete()
    return redirect('decoupliste')

# NIVEAUX------------------------------------------------------------------------------------------------------------
# liste 
@login_required()
def nivliste(request):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    niv = Niveaux.objects.all()
    return render(request,'Pages/niveaux/liste.html', context={'niv':niv,'annee': statuta})


def nivlistecat(request):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    niv = Niveaux.objects.all()
    for ni in niv:
        cat = Catalogues.objects.filter(Niveau=ni.id) # type: ignore
    
    return render(request, 'Pages/niveaux/listecat.html', context={'niv': niv, 'cat': cat,'annee': statuta}) # type: ignore

# ajouter
@login_required()
def nivajout(request):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    if request.method == 'POST':
        fm = NiveauRegistration(request.POST)
        if fm.is_valid():
            fm.save()
            messages.success(request, 'Votre formulaire a été envoyé avec succès !')
            fm = NiveauRegistration()
            return redirect('nivliste')
    else :
        fm = NiveauRegistration()
    return render(request,'Pages/niveaux/ajout.html',{'form':fm,'annee': statuta}  )

# modifier
@login_required()
def nivmodif(request, id):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    if request.method == 'POST':
        pi = Niveaux.objects.get(pk=id)
        fm = NiveauRegistration(request.POST, instance=pi)
        if fm.is_valid():
            fm.save()
            return redirect('nivliste')
    else:
        pi = Niveaux.objects.get(pk=id)
        fm = NiveauRegistration(instance=pi)
    return render(request, 'Pages/niveaux/modif.html',{'form':fm,'annee': statuta})

# supprimer
@login_required()
def delete_niv(request, id):
    niv = Niveaux.objects.get(pk=id)
    niv.delete()
    return redirect('nivliste')


# HISTORIQUES ABSENCES------------------------------------------------------------------------------------------------------------
# liste 
@login_required()
def absliste(request):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    abs = HistoriquesAbs.objects.all()
    return render(request,'Pages/heuresabsences/liste.html', context={'abs':abs,'annee': statuta})


# ajouter
@login_required()
def absajout(request,object_id):
    
    statuta = Anneescolaires.objects.get(Statutannee='1')
    statuti = Decoupages.objects.get(Statutdecoup='1')
    student = Etudiants.objects.get(id=object_id)
    dossier_scolaire = DossierScolaires.objects.get(Etudiant=student, Anneescolaire=statuta)
    if request.method == 'POST':
        fm = HistoriqueabsRegistration(request.POST)
        if fm.is_valid():

            cours = fm.cleaned_data['Cours']
            heuredebut = fm.cleaned_data['Heuredebut']
            heurefin = fm.cleaned_data['Heurefin']
            histo = HistoriquesAbs.objects.create(Decoupage=statuti,Anneescolaire=statuta, Etudiant = student,Cours = cours, Heuredebut= heuredebut,Heurefin=heurefin)
            
            # Convertir les objets time en objets datetime
            now = datetime.now()
            debut = datetime.combine(now.date(), heuredebut)
            fin = datetime.combine(now.date(), heurefin)

            # Calculer la durée entre les horaires
            duree = fin - debut
            duree_en_secondes = duree.total_seconds()/ 3600

            # Modifier les heures d'absence correspondantes dans le dossier scolaire
            if statuti.Decoupage == 'Semestre 1':
                dossier_scolaire.HeuresabsenceS1 += int(duree_en_secondes)
            elif statuti.Decoupage == 'Semestre 2':
                dossier_scolaire.HeuresabsenceS2 += int(duree_en_secondes)

            dossier_scolaire.save()

            messages.success(request, 'Votre formulaire a été envoyé avec succès !')
            fm = HistoriqueabsRegistration()
            return redirect('absliste')
    else :
        fm = HistoriqueabsRegistration()
    return render(request,'Pages/heuresabsences/ajout.html',{'form':fm,'annee': statuta}  )

# modifier
@login_required()
def absmodif(request, id):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    if request.method == 'POST':
        pi = HistoriquesAbs.objects.get(pk=id)
        fm = HistoriqueabsRegistration(request.POST, instance=pi)
        if fm.is_valid():
            fm.save()
            return redirect('absliste')
    else:
        pi = HistoriquesAbs.objects.get(pk=id)
        fm = HistoriqueabsRegistration(instance=pi)
    return render(request, 'Pages/heuresabsences/modif.html',{'form':fm,'annee': statuta})

# supprimer
@login_required()
def delete_abs(request, id):
    abs = HistoriquesAbs.objects.get(pk=id)
    abs.delete()
    return redirect('absliste')

# RAPPORT------------------------------------------------------------------------------------------------------------
def generer_rapport(request):
    # Création du document Word
    document = docx.Document()

    # Ajouter un titre au rapport
    titre = "Rapport de synthèse de gestion de l'université"
    document.add_heading(titre, level=1)

    # Récupérer les informations de l'université depuis la base de données
    universite = Salles.objects.get(id=1)
    annee = universite.Salle
    nombre_etudiants = universite.Nbreplace
    nombre_enseignants = universite.Professeurprincipal
    budget_total = universite.Planning

    # Ajouter les informations au rapport
    document.add_paragraph(f"Université : {universite.Salle}")
    document.add_paragraph(f"Année : {annee}")
    document.add_paragraph(f"Nombre d'étudiants : {nombre_etudiants}")
    document.add_paragraph(f"Nombre d'enseignants : {nombre_enseignants}")
    document.add_paragraph(f"Budget total : {budget_total}")

    # Enregistrer le document Word
    document.save("rapport_synthese.docx")

    # Renvoyer le fichier en tant que réponse HTTP
    file_path = "rapport_synthese.docx"
    return FileResponse(open(file_path, 'rb'), as_attachment=True)

# CATALOGUES------------------------------------------------------------------------------------------------------------
# ajouter
@login_required()
def add_catalogue(request,id):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    salle = Salles.objects.get(id=id)
    Niv = Niveaux.objects.get(id=id)
    if request.method == 'POST':
        
        fm = CatalogueRegistration(request.POST)
        if fm.is_valid():
            nom = fm.cleaned_data['Nom']
            prix = fm.cleaned_data['Prix']
            serm = Catalogues.objects.create(Nom=nom,Prix=prix, Niveau=Niv)

          
            messages.success(request, 'Votre formulaire a été envoyé avec succès !')
            fm = SessionRegistration()
            return redirect('nivlistecat')
    else :
        fm = CatalogueRegistration()
    return render(request,'Pages/catalogues/ajout.html',{'form':fm,'annee': statuta}  )

# DOSSIER SCOLAIRES------------------------------------------------------------------------------------------------------------
# liste 
@login_required()
def dossierliste(request,id):
    etudian = Etudiants.objects.get(id=id)
    statuta = Anneescolaires.objects.get(Statutannee='1')
    salle = Salles.objects.filter(Salle=etudian.Salle,Etablissement=etudian.Ecole, Anneescolaire=statuta).get()
    dat = Documents.objects.filter(Salle=salle).order_by('Typematiere')
    etudia = Etudiants.objects.get(id=id)
    etudiant = Etudiants.objects.get(id=id)
    sessio = Decoupages.objects.filter(id='2')
    sessi = Decoupages.objects.filter(id='1')
    sessi1 = Decoupages.objects.get(id='1')
    sessio1 = Decoupages.objects.get(id='2')
    parent=Contacturgences.objects.get(Etudiant=id)
    
    sall = Salles.objects.filter(Salle=etudian.Salle, Anneescolaire=statuta)

    grades = Notes.objects.filter(Etudiant=etudian, Anneescolaire=statuta, Decoupage=sessi1)
    grade = Notes.objects.filter(Etudiant=etudian, Anneescolaire=statuta, Decoupage=sessio1)

    
    statuti = Decoupages.objects.get(Statutdecoup='1')

    abs = HistoriquesAbs.objects.filter(Etudiant=etudiant, Anneescolaire=statuta, Decoupage__Decoupage='Semestre 1')
    abs1 = HistoriquesAbs.objects.filter(Etudiant=etudiant, Anneescolaire=statuta, Decoupage__Decoupage='Semestre 2')
   
    context={'data':dat,'annee': statuta,'etudiant':etudiant,'sessi':sessi,'sessio':sessio,'grades':grades,'grade':grade,'sall':sall,'annee':statuta,'parent':parent,'abs':abs,'abs1':abs1}
    return render(request,'Pages/dossierscolaires/liste.html', context)

# PERSONNELS........................................................................................................................................
# liste 
@login_required()
def personnelliste(request):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    person = Personnels.objects.all()
    return render(request,'Pages/personnels/liste.html', context={'person':person,'annee': statuta})

# liste 
@login_required()
def personnellistecont(request):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    person = Personnels.objects.all()
    return render(request,'Pages/personnels/listecontact.html', context={'person':person,'annee': statuta})

# ajouter
@login_required()
def personnelajout(request):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    if request.method == 'POST':
        fm = PersonnelRegistration(request.POST, request.FILES)
        if fm.is_valid():
            Photo = fm.cleaned_data['Photo']
            
            student = fm.save(commit=False)
           

            student.save()
            # Créer un utilisateur avec le même nom d'utilisateur et mot de passe que l'étudiant
            userna = fm.cleaned_data['Nom']
            usernam = fm.cleaned_data['Prenoms']
            FString = userna + " " + usernam
            username = FString
            role = "ACADEMIQUE"
            statu = 'NON ACTIVE'
            password = 'P@ssword'
            profile_photo = Photo
            encoded_password = make_password(password)
            email = fm.cleaned_data['Email']
            userm = Utilisateur.objects.create(statut=statu,profile_photo=profile_photo,username=username,last_name=usernam,first_name=userna, password=encoded_password, email=email,role=role)

            # Ajouter l'utilisateur à l'objet Etudiant et enregistrer à nouveau
            student.User = userm
            student.save()
            return redirect('listepersonnelrh')
    else :
        matricule = str(uuid.uuid4())[:8] # génération d'une chaîne aléatoire de 8 caractères
        fm = PersonnelRegistration(initial={'Matricule': matricule,'Statut': "NON ACTIVE"})
    return render(request,'Pages/personnels/ajout.html',{'form':fm,'annee': statuta}  )

# AUTORISATIONS........................................................................................................................................
# liste personnel
@login_required()
def personnelauto(request):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    person = Autorisationspersonnel.objects.all()
    return render(request,'Pages/autorisationspersonnel/liste.html', context={'person':person,'annee': statuta})

@login_required
def ajoutautopers(request, id, iid, iiid):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    if request.method == 'POST':
        fm = AutopersRegistration(request.POST)
        person = Personnels.objects.get(id=id)
        perso = person.Service
        pers = person.Poste

        if fm.is_valid():
            autopers = fm.save(commit=False)
            autopers.save()
            return redirect('listepersonnelauto')
    else:
        fm = AutopersRegistration(initial={'Personnel': id, 'Service': iid, 'Poste': iiid})
    
    context = {
        'form': fm,
        'id': id,
        'iid': iid,
        'iiid': iiid,'annee': statuta
    }
    return render(request, 'Pages/autorisationspersonnel/ajout.html', context)

# modifier
def autopersmodif(request, id):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    if request.method == 'POST':
        pi = Autorisationspersonnel.objects.get(pk=id)
        fm = AutopersRegistration(request.POST, instance=pi)
        if fm.is_valid():
            fm.save()
            return redirect('listepersonnelauto')
    else:
        pi = Autorisationspersonnel.objects.get(pk=id)
        fm = AutopersRegistration(instance=pi)
    return render(request, 'Pages/autorisationspersonnel/modif.html',{'form':fm,'annee': statuta})

# supprimer
@login_required()
def autopersdelete(request, id):
    autopers = Autorisationspersonnel.objects.get(pk=id)
    autopers.delete()
    return redirect('listepersonnelauto')

# AUTORISATIONS ENSEIGNANT........................................................................................................................................
# liste personnel
@login_required()
def ensauto(request):
    person = Autorisationsenseignant.objects.all()
    statuta = Anneescolaires.objects.get(Statutannee='1')
    return render(request,'Pages/autorisationsenseignants/liste.html', context={'person':person,'annee': statuta})

@login_required
def ajoutautoens(request, id, iid):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    if request.method == 'POST':
        fm = AutoensRegistration(request.POST)
        if fm.is_valid():
            autopers = fm.save(commit=False)
            autopers.save()
            return redirect('listeensauto')
    else:
        fm = AutoensRegistration(initial={'Enseignant': id, 'Ecole': iid})
    
    context = {
        'form': fm,
        'id': id,
        'iid': iid,'annee': statuta
    }
    return render(request, 'Pages/autorisationsenseignants/ajout.html', context)

# modifier
def autoensmodif(request, id):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    if request.method == 'POST':
        pi = Autorisationsenseignant.objects.get(pk=id)
        fm = AutoensRegistration(request.POST, instance=pi)
        if fm.is_valid():
            fm.save()
            return redirect('listeensauto')
    else:
        pi = Autorisationsenseignant.objects.get(pk=id)
        fm = AutoensRegistration(instance=pi)
    return render(request, 'Pages/autorisationsenseignants/modif.html',{'form':fm,'annee': statuta})

# supprimer
@login_required()
def autoensdelete(request, id):
    autopers = Autorisationsenseignant.objects.get(pk=id)
    autopers.delete()
    return redirect('listeensauto')


# AUTORISATIONS ENSEIGNANT........................................................................................................................................
# liste personnel
@login_required()
def audauto(request):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    person = Autorisationsauditeur.objects.all()
    return render(request,'Pages/autorisationsauditeurs/liste.html', context={'person':person,'annee': statuta})

@login_required
def ajoutautoaud(request, id, iid):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    student = Personnels.objects.get(id=id)
    if request.method == 'POST':
        fm = AutoaudRegistration(request.POST)
        if fm.is_valid():
            autopers = fm.save(commit=False)
            autopers.save()
            return redirect('listeaudauto')
    else:
        fm = AutoaudRegistration(initial={'Etudiant': id, 'Salle': iid})
    
    context = {
        'form': fm,
        'id': id,
        'iid': iid,
        'student_name': student.Nom,
        'student_surname': student.Prenoms,'annee': statuta
    }
    return render(request, 'Pages/autorisationsauditeurs/ajout.html', context)

# modifier
def autoaudmodif(request, id):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    if request.method == 'POST':
        pi = Autorisationsauditeur.objects.get(pk=id)
        fm = AutoaudRegistration(request.POST, instance=pi)
        if fm.is_valid():
            fm.save()
            return redirect('listeensauto')
    else:
        pi = Autorisationsauditeur.objects.get(pk=id)
        fm = AutoaudRegistration(instance=pi)
    return render(request, 'Pages/autorisationsauditeurs/modif.html',{'form':fm,'annee': statuta})

# supprimer
@login_required()
def autoauddelete(request, id):
    autopers = Autorisationsauditeur.objects.get(pk=id)
    autopers.delete()
    return redirect('listeaudauto')


# MISSIONS........................................................................................................................................
# liste personnel
@login_required()
def mispersonnel(request):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    person = Missions.objects.all()
    return render(request,'Pages/missions/liste.html', context={'person':person,'annee': statuta})

@login_required
def misajout(request, id, iid, iiid):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    if request.method == 'POST':
        fm = MisRegistration(request.POST)
        person = Personnels.objects.get(id=id)
        perso = person.Service
        pers = person.Poste

        if fm.is_valid():
            autopers = fm.save(commit=False)
            autopers.save()
            return redirect('listemis')
    else:
        fm = MisRegistration(initial={'Personnel': id, 'Service': iid, 'Poste': iiid})
    
    context = {
        'form': fm,
        'id': id,
        'iid': iid,
        'iiid': iiid,'annee': statuta
    }
    return render(request, 'Pages/missions/ajout.html', context)

# modifier
def mismodif(request, id):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    if request.method == 'POST':
        pi = Missions.objects.get(pk=id)
        fm = MisRegistration(request.POST, instance=pi)
        if fm.is_valid():
            fm.save()
            return redirect('listemis')
    else:
        pi = Missions.objects.get(pk=id)
        fm = MisRegistration(instance=pi)
    return render(request, 'Pages/missions/modif.html',{'form':fm,'annee': statuta})

# supprimer
@login_required()
def misdelete(request, id):
    autopers = Missions.objects.get(pk=id)
    autopers.delete()
    return redirect('listemis')


def get(self, request):
        result = (Personnels.objects.all()
                  .values('Nom', 'Prenoms')
                  .annotate(type='personnels')
                  .filter(Q(Service__isnull=False) | Q(Poste__isnull=False))
                  .union(Etudiants.objects.all()
                         .values('Nom', 'Prenoms')
                         .annotate(type='etudiants')
                         .filter(Q(Ecole__isnull=False) | Q(Niveau__isnull=False) | Q(Salle__isnull=False)))
                  .union(Enseignants.objects.all()
                         .values('Nom', 'Prenoms')
                         .annotate(type='enseignants')
                         .filter(Q(Ecole__isnull=False) | Q(Statut__isnull=False)))
                  .order_by('Nom', 'Prenoms')
                 )
        
        context = {
            'result': result
        }
        
        return render(request, 'reunion_tables.html', context)

# FINANCES........................................................................................................................................
# liste etudiant non inscrit
@login_required(login_url='/accounts/')

def etunon(request):
    annee_scolaire_active = Anneescolaires.objects.get(Statutannee='1')

    etudiants_avec_dossiers_scolaires = Etudiants.objects.filter(
        dossiers_scolaires__Anneescolaire=annee_scolaire_active
    ).distinct()

    etudiants_avec_dossiers_finances = Etudiants.objects.filter(
        dossiers_finances__Anneescolaire=annee_scolaire_active
    ).distinct()

    etudiants_sans_dossiers_finances = etudiants_avec_dossiers_scolaires.exclude(
        Q(id__in=etudiants_avec_dossiers_finances)
    )
    etud = etudiants_sans_dossiers_finances
    return render(request,'Pages/etudiants/listenoninscrit.html', context={'etu':etud,'annee': annee_scolaire_active})

# Ajout de paiement inscription
def historiquevers(request, niveau_id, etudiant_id):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    niveau = get_object_or_404(Niveaux, id=niveau_id)
    etudiant = get_object_or_404(Etudiants, id=etudiant_id)
    catalogues = Catalogues.objects.filter(Niveau=niveau)
    
    return render(request, 'Pages/Historiquesversements/ajout.html', {'niveau': niveau, 'etudiant': etudiant, 'catalogues': catalogues,'annee': statuta})

def savevers(request):
    if request.method == 'POST':
        niveau_id = request.POST.get('niveau_id')
        etudiant_id = request.POST.get('etudiant_id')
        montantvers = request.POST.get('montantvers')
        
        niveau = get_object_or_404(Niveaux, id=niveau_id)
        etudiant = get_object_or_404(Etudiants, id=etudiant_id)
        statuta = Anneescolaires.objects.get(Statutannee='1')
        catalogues = Catalogues.objects.filter(Niveau=niveau).get()
        Ecole = Ecoles.objects.filter(Etablissement=etudiant.Ecole).get()
        
        
        # Récupérer les autres informations à partir de l'URL ou des données du formulaire
        dos = DossierScolaires.objects.filter(Etudiant=etudiant_id, Anneescolaire=statuta).get()
        anneescolaire = statuta
        datevers = request.POST.get('datevers')
        montantscolarite = catalogues.Prix
        montantrestant = '0'
        nom = dos.Nom
        ecole = dos.Ecole
        sexe = dos.Sexe
        salle = dos.Salle
        
        # Vérifier si les valeurs numériques ne sont pas None
        if montantscolarite is not None:
            montantscolarite = float(montantscolarite)
        if montantrestant is not None:
            montantrestant = float(montantrestant)
        if montantvers is not None:
            montantvers = float(montantvers)
        
        dossierfinance, created = DossierFinances.objects.get_or_create(Etudiant=etudiant, Anneescolaire=statuta)

        if created:
            dossierfinance.Montantscolarit = montantscolarite # type: ignore
            dossierfinance.Montantvers = 0
            dossierfinance.Montantrestant = montantscolarite # type: ignore
            dossierfinance.Matricule = etudiant.Matricule
            dossierfinance.Nom = etudiant.Nom
            dossierfinance.Ecole = ecole
            dossierfinance.Sexe = sexe
            dossierfinance.Niveau = niveau
            dossierfinance.Salle = salle
            dossierfinance.save()

        historique_vers = HistoriquesVers.objects.create(
            Etudiant=etudiant,
            Niveau=niveau,
            Montantvers=montantvers,
            Anneescolaire=anneescolaire,
            Datevers=date.today(),
            Montantscolarit=montantscolarite,
            Montantrestant=montantscolarite-montantvers,# type: ignore
            Nom=nom,
            Ecole=ecole,
            Sexe=sexe,
            Salle=salle,
            Dossierfinance=dossierfinance,
        )
        
        # Enregistrer l'historique de versement
        historique_vers.save()
        
        dossier_finances = DossierFinances.objects.get_or_create(Etudiant=etudiant, Anneescolaire=statuta)[0]
        if montantvers is not None:
            dossier_finances.Montantvers = dossier_finances.Montantvers or 0
            dossier_finances.Montantvers += int(montantvers)
            dossier_finances.Montantrestant = dossier_finances.Montantscolarit or 0
            dossier_finances.Montantrestant -= int(montantvers)
        dossier_finances.save()      
        
        
        # Rediriger vers une autre vue ou une page de succès
        return redirect('listenon')

# liste etudiant inscrit
def listeinscrit(request):
    active_year = Anneescolaires.objects.get(Statutannee='1')
    students = Etudiants.objects.filter(
        Q(dossiers_finances__Anneescolaire=active_year)
    )
    etud = students
    return render(request,'Pages/etudiants/listeinscrit.html', context={'etu':etud,'annee': active_year})

# Ajout de versement
def historiqueversautre(request, niveau_id, etudiant_id):
    annee_scolaire_active = Anneescolaires.objects.get(Statutannee='1')
    niveau = get_object_or_404(Niveaux, id=niveau_id)
    etudiant = get_object_or_404(Etudiants, id=etudiant_id)
    catalogues = Catalogues.objects.filter(Niveau=niveau)
    statuta = Anneescolaires.objects.get(Statutannee='1')
    dos = DossierFinances.objects.filter(Etudiant=etudiant_id, Anneescolaire=statuta) 

    return render(request, 'Pages/Historiquesversements/ajoutvers.html', {'dos': dos, 'niveau': niveau, 'etudiant': etudiant, 'catalogues': catalogues,'annee': annee_scolaire_active})


def saveversautre(request):
    if request.method == 'POST':
        niveau_id = request.POST.get('niveau_id')
        etudiant_id = request.POST.get('etudiant_id')
        montantvers = request.POST.get('montantvers')

        niveau = get_object_or_404(Niveaux, id=niveau_id)
        etudiant = get_object_or_404(Etudiants, id=etudiant_id)
        statuta = Anneescolaires.objects.get(Statutannee='1')
        catalogues = Catalogues.objects.filter(Niveau=niveau).get()
        ecole = Ecoles.objects.get(Etablissement=etudiant.Ecole)

        # Récupérer les autres informations à partir de l'URL ou des données du formulaire
        dos = DossierScolaires.objects.get(Etudiant=etudiant_id, Anneescolaire=statuta)
        doss = DossierFinances.objects.filter(Etudiant=etudiant_id, Anneescolaire=statuta).get()
        anneescolaire = statuta
        datevers = date.today()
        montantscolarite = catalogues.Prix
        montantrestant1 = doss.Montantrestant
        nom = dos.Nom
        ecole = dos.Ecole
        sexe = dos.Sexe
        salle = dos.Salle

        dossierfinance = DossierFinances.objects.get(Etudiant=etudiant_id, Anneescolaire=statuta)

        # Vérifier si les valeurs numériques ne sont pas None
        if montantscolarite is not None:
            montantscolarite = int(montantscolarite)
        if montantvers is not None:
            montantvers = int(montantvers)

        historique_vers = HistoriquesVers.objects.create(
            Etudiant=etudiant,
            Niveau=niveau,
            Montantvers=montantvers,
            Anneescolaire=anneescolaire,
            Datevers=date.today(),
            Montantscolarit=montantscolarite,
            Montantrestant=montantrestant1 - montantvers,# type: ignore
            Nom=nom,
            Ecole=ecole,
            Sexe=sexe,
            Salle=salle,
            Dossierfinance=dossierfinance,
        )

        # Enregistrer l'historique de versement
        historique_vers.save()

        dossierfinance.Montantvers += montantvers# type: ignore
        dossierfinance.Montantrestant -= montantvers# type: ignore
        dossierfinance.save()

        # Rediriger vers une autre vue ou une page de succès
        return redirect('listeinscrit')

# liste paiement
def listepaiejour(request):
    aujourdhui = date.today()
    statuta = Anneescolaires.objects.get(Statutannee='1')
    # Effectuer la requête pour récupérer les HistoriquesVers de la date du jour
    historiquesvers = HistoriquesVers.objects.filter(Datevers=aujourdhui)
    etud = historiquesvers
    return render(request,'Pages/Historiquesversements/listejour.html', context={'etu':etud,'annee': statuta})

def listepaieannee(request):
    statuta = Anneescolaires.objects.get(Statutannee='1')

    # Effectuer la requête pour récupérer les HistoriquesVers de la date du jour
    historiquesvers = HistoriquesVers.objects.filter(Anneescolaire=statuta)
    etud = historiquesvers
    return render(request,'Pages/Historiquesversements/listeannee.html', context={'etu':etud,'annee': statuta})

# TYPES AVIS------------------------------------------------------------------------------------------------------------
# liste 
@login_required()
def avisliste(request):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    avis = Avis.objects.all()
    return render(request,'Pages/avis/liste.html', context={'avis':avis,'annee': statuta})

# ajouter
@login_required()
def avisajout(request):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    if request.method == 'POST':
        fm = AvisRegistration(request.POST)
        if fm.is_valid():
            
            fm.save()
            messages.success(request, 'Votre formulaire a été envoyé avec succès !')
            fm = AvisRegistration()
            return redirect('avisliste')
    else :
        fm = AvisRegistration()
    return render(request,'Pages/avis/ajout.html',{'form':fm,'annee': statuta}  )

# modifier
@login_required()
def avismodif(request, id):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    if request.method == 'POST':
        pi = Avis.objects.get(pk=id)
        fm = AvisRegistration(request.POST, instance=pi)
        if fm.is_valid():
            fm.save()
            return redirect('avisliste')
    else:
        pi = Avis.objects.get(pk=id)
        fm = AvisRegistration(instance=pi)
    return render(request, 'Pages/avis/modif.html',{'form':fm,'annee': statuta})

# MATRIÙONIALES------------------------------------------------------------------------------------------------------------
# liste 
@login_required()
def matriliste(request):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    avis = Avis.objects.all()
    return render(request,'Pages/avis/liste.html', context={'avis':avis,'annee': statuta})

# ajouter
@login_required()
def matriajout(request):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    if request.method == 'POST':
        fm = AvisRegistration(request.POST)
        if fm.is_valid():
            
            fm.save()
            messages.success(request, 'Votre formulaire a été envoyé avec succès !')
            fm = AvisRegistration()
            return redirect('avisliste')
    else :
        fm = AvisRegistration()
    return render(request,'Pages/avis/ajout.html',{'form':fm,'annee': statuta}  )

# modifier
@login_required()
def matrimodif(request, id):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    if request.method == 'POST':
        pi = Avis.objects.get(pk=id)
        fm = AvisRegistration(request.POST, instance=pi)
        if fm.is_valid():
            fm.save()
            return redirect('avisliste')
    else:
        pi = Avis.objects.get(pk=id)
        fm = AvisRegistration(instance=pi)
    return render(request, 'Pages/avis/modif.html',{'form':fm,'annee': statuta})

# TYPES AVIS -----------------------------------------------------------------------------------------------------------------------------------------------------------------
# supprimer
@login_required()
def avisdelete(request, id):
    avis = Avis.objects.get(pk=id)
    avis.delete()
    return redirect('avisliste')

# USERS
# liste 
@login_required()
def user(request):
    statuta = Anneescolaires.objects.get(Statutannee='1')
    enseignant = Utilisateur.objects.all()
    return render(request,'Pages/users/listeutilisateur.html', context={'enseignantliste':enseignant,'annee': statuta})

